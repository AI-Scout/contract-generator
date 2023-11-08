import streamlit as st
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from docx import Document
import os


st.header("Contract Generator Demo")

openai_api_key = st.secrets["OPEN_AI_API_KEY"]

st.sidebar.subheader('About')
st.sidebar.markdown('<b>Created by [AI Scout Solutions](https://solutions.aiscout.net/)</b>. For demo purposes only- this is not the final product.', unsafe_allow_html=True)

st.sidebar.subheader('Purpose')
st.sidebar.markdown('Generates a completed version of this [REAL ESTATE LISTING AGREEMENT](https://docdro.id/EQpSLuC) with AI, based on user response.', unsafe_allow_html=True)

st.sidebar.subheader('How it Works')
st.sidebar.caption('1. Answer all questions.')
st.sidebar.caption('2. AI will generate the contract.')
st.sidebar.caption('3. You may download the completed contract as a Word Doc.')

file_path = None

def load_template(filename):
    with open(filename, 'r') as f:
        return f.read()

def process_form_data(seller_address, agency_address, listing_rights, purchase_price, 
                      listing_period, commission_due, commission_type, commission_amount, 
                      authorize_advertise, third_party_estimation, other_offers_disclosure, 
                      lockbox_auth, for_sale_sign, media_permission, dual_agency, 
                      property_disclosure, agency_other_clients, binding_arbitration, 
                      governing_state, additional_terms, received_signed_copy, 
                      date_of_agreement, seller_name, broker_name, agency_name, 
                      address_of_real_property, tax_map_lot, deed_book_page, 
                      other_info_about_property, fixtures_not_to_sell, personal_property_included, 
                      leasing_percentage, deed_type, other_agents_comission):    

        with st.spinner("Generating contract..."):
            # Instantiate LLM model
            llm = OpenAI(model_name="gpt-3.5-turbo-16k", temperature=0.33, openai_api_key=openai_api_key)
            
            # Load the template and format it
            template = load_template('prompt.txt')
            prompt_query = template.format(
                seller_address=seller_address,
                agency_address=agency_address,
                listing_rights=listing_rights,
                purchase_price=purchase_price,
                listing_period=listing_period,
                commission_due=commission_due,
                commission_type=commission_type,
                commission_amount=commission_amount,
                authorize_advertise=authorize_advertise,
                third_party_estimation=third_party_estimation,
                other_offers_disclosure=other_offers_disclosure,
                lockbox_auth=lockbox_auth,
                for_sale_sign=for_sale_sign,
                media_permission=media_permission,
                dual_agency=dual_agency,
                property_disclosure=property_disclosure,
                agency_other_clients=agency_other_clients,
                binding_arbitration=binding_arbitration,
                governing_state=governing_state,
                additional_terms=additional_terms,
                received_signed_copy=received_signed_copy,
                date_of_agreement=date_of_agreement,
                seller_name=seller_name,
                broker_name=broker_name,
                agency_name=agency_name,
                address_of_real_property=address_of_real_property,
                tax_map_lot=tax_map_lot,
                deed_book_page=deed_book_page,
                other_info_about_property=other_info_about_property,
                fixtures_not_to_sell=fixtures_not_to_sell,
                personal_property_included=personal_property_included,
                leasing_percentage=leasing_percentage,
                deed_type=deed_type,
                other_agents_comission=other_agents_comission
            )
    
            # Run LLM model
            response = llm(prompt_query)
            
            lines = response.split('\r\n')
            
            # Create Word Document
            doc = Document()
            doc.add_heading('Contract', level=1)
        
        # Add each line of the response as a separate paragraph
        for line in lines:
            doc.add_paragraph(line.strip())  # .strip() removes leading and trailing whitespaces
            
        # Save the Word Document to a temporary file
        temp_file = "temp_contract.docx"
        doc.save(temp_file)
        
        return temp_file

with st.form("listing_form"):
    seller_address = st.text_input("What is the mailing address of the seller?", 
                                   placeholder="Ex: 123 Elm Street, Springfield, IL 62704")
    agency_address = st.text_input("What is the mailing address of the agency?", 
                                   placeholder="Ex: 456 Oak Avenue, Springfield, IL 62701")
    listing_rights = st.selectbox("Do you grant the agency exclusive right-to-sell, exclusive agency, or open listing rights?", ["Exclusive right-to-sell", "Exclusive agency", "Open listing"])
    date_of_agreement = st.date_input("What is the date of the agreement?")
    seller_name = st.text_input("What is the name of the seller?")
    broker_name = st.text_input("What is the name of the broker?")
    agency_name = st.text_input("What is the name of the agency?")
    address_of_real_property = st.text_input("What is the address of the real property?")
    tax_map_lot = st.text_input("What is the Tax Map/Lot of the property?")
    deed_book_page = st.text_input("What is the deed book/page of the property?")
    other_info_about_property = st.text_area("Any other info about the property?")
    fixtures_not_to_sell = st.text_input("What fixtures should not be sold?")
    personal_property_included = st.text_input("What personal property shall be included?")
    leasing_percentage = st.text_input("What is the leasing percentage?")
    deed_type = st.text_input("What type of deed?")
    purchase_price = st.text_input("What is the purchase price for the property?", 
                                   placeholder="Ex: $375,000")
    listing_period = st.text_input("When does the listing agreement start and end?", 
                                   placeholder="Ex: September 25, 2023 to March 25, 2024")
    commission_due = st.selectbox("If the property is sold within a certain period after the expiration of the listing period, are any commission fees due?", ["Yes", "No"])
    commission_type = st.selectbox("Do you prefer a percentage commission or a fixed payment commission?", ["Percentage commission", "Fixed payment commission"])
    commission_amount = st.text_input("What is the commission amount? (Provide as a percentage or fixed value based on your choice above)")
    other_agents = st.selectbox("Are there any other licensed real estate agents?", ["Yes", "No"])
    other_agents_comission = st.text_input("If yes, what is their commission?") if other_agents == "Yes" else None
    authorize_advertise = st.selectbox("Are you authorizing the agency to advertise the property and use its street address in advertising?", ["Yes", "No"])
    third_party_estimation = st.selectbox("Do you authorize third-party websites to create estimated market values of the property?", ["Yes", "No"])
    other_offers_disclosure = st.selectbox("Are you allowing the agency to disclose the existence of other offers on the property?", ["Yes", "No"])
    lockbox_auth = st.selectbox("Are you authorizing the placement of a lockbox or key box on the property?", ["Yes", "No"])
    for_sale_sign = st.selectbox("Do you authorize a 'For Sale' sign on the property?", ["Yes", "No"])
    media_permission = st.selectbox("Are you giving permission for the agency to take and use photos and videos of the property for marketing purposes?", ["Yes", "No"])
    dual_agency = st.selectbox("Are you allowing disclosed dual agency, or not allowing it?", ["Allowing", "Not Allowing"])
    property_disclosure = st.selectbox("Have you reviewed and completed the Property Disclosure Statement accurately?", ["Yes", "No"])
    agency_other_clients = st.selectbox("Do you acknowledge that the agency may have other clients with similar properties?", ["Yes", "No"])
    binding_arbitration = st.selectbox("Are you agreeing to binding arbitration in case of disputes?", ["Yes", "No"])
    governing_state = st.text_input("What state's laws will govern this agreement?", 
                                   placeholder="Ex: Illinois")
    additional_terms = st.text_area("Are there any additional terms and conditions you want to include?")
    received_signed_copy = st.selectbox("Have you received a signed copy of this agreement?", ["Yes", "No"])

    submitted = st.form_submit_button("Submit")
    
    if submitted:
        file_path = process_form_data(
            seller_address=seller_address,
            agency_address=agency_address,
            listing_rights=listing_rights,
            purchase_price=purchase_price,
            listing_period=listing_period,
            commission_due=commission_due,
            commission_type=commission_type,
            commission_amount=commission_amount,
            authorize_advertise=authorize_advertise,
            third_party_estimation=third_party_estimation,
            other_offers_disclosure=other_offers_disclosure,
            lockbox_auth=lockbox_auth,
            for_sale_sign=for_sale_sign,
            media_permission=media_permission,
            dual_agency=dual_agency,
            property_disclosure=property_disclosure,
            agency_other_clients=agency_other_clients,
            binding_arbitration=binding_arbitration,
            governing_state=governing_state,
            additional_terms=additional_terms,
            received_signed_copy=received_signed_copy,
            date_of_agreement=date_of_agreement,
            seller_name=seller_name,
            broker_name=broker_name,
            agency_name=agency_name,
            address_of_real_property=address_of_real_property,
            tax_map_lot=tax_map_lot,
            deed_book_page=deed_book_page,
            other_info_about_property=other_info_about_property,
            fixtures_not_to_sell=fixtures_not_to_sell,
            personal_property_included=personal_property_included,
            leasing_percentage=leasing_percentage,
            deed_type=deed_type,
            other_agents_comission=other_agents_comission
        )

if file_path:
    with open(file_path, "rb") as file:
        st.download_button(
            label="Download Contract",
            data=file,
            file_name="contract.docx",
            mime="application/msword",
        )
    # Remove the temporary file
    os.remove(file_path)

hide_streamlit_style = """    
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 
